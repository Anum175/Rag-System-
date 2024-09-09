import streamlit as st
from llama_index.core import VectorStoreIndex, Document
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.ollama import Ollama
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

st.title("RAG SYSTEM")

# Set up the embedding model and LLM
embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-base-en-v1.5")
llm = Ollama(model="llama2", request_timeout=360.0)

# Initialize session state
if "chats" not in st.session_state:
    st.session_state.chats = {"General": {"messages": [], "documents": [], "index": None}}
if "current_chat" not in st.session_state:
    st.session_state.current_chat = "General"
if "edit_mode" not in st.session_state:
    st.session_state.edit_mode = {}
if "new_chat_name" not in st.session_state:
    st.session_state.new_chat_name = ""

# Sidebar for chat management
st.sidebar.title("Chat Management")

# Add new chat
new_chat_name = st.sidebar.text_input("Enter new chat name:", key="new_chat_name_input")
if st.sidebar.button("Add New Chat"):
    if new_chat_name and new_chat_name not in st.session_state.chats:
        st.session_state.chats[new_chat_name] = {"messages": [], "documents": [], "index": None}
        st.session_state.current_chat = new_chat_name
        st.session_state.new_chat_name = ""
        st.session_state.edit_mode = {}
        st.experimental_rerun()
    elif new_chat_name in st.session_state.chats:
        st.sidebar.error("Chat already exists.")
    else:
        st.sidebar.error("Please enter a chat name.")

# Display chats with options
for chat_name in list(st.session_state.chats.keys()):
    col1, col2, col3 = st.sidebar.columns([3, 1, 1])
    
    if col1.button(chat_name, key=f"select_{chat_name}"):
        st.session_state.current_chat = chat_name
        st.experimental_rerun()
    
    if col2.button("✏️", key=f"edit_{chat_name}"):
        st.session_state.edit_mode[chat_name] = True
        st.session_state.new_chat_name = chat_name
    
    if col3.button("🗑️", key=f"delete_{chat_name}"):
        if chat_name != "General":
            del st.session_state.chats[chat_name]
            if st.session_state.current_chat == chat_name:
                st.session_state.current_chat = "General"
            st.session_state.edit_mode.pop(chat_name, None)
            st.session_state.new_chat_name = ""
            st.experimental_rerun()
        else:
            st.sidebar.error("Cannot delete General chat.")

    if st.session_state.edit_mode.get(chat_name, False):
        new_name = st.sidebar.text_input("New name:", value=chat_name, key=f"edit_input_{chat_name}")
        if st.sidebar.button("Save", key=f"save_{chat_name}"):
            if new_name and new_name != chat_name:
                st.session_state.chats[new_name] = st.session_state.chats.pop(chat_name)
                if st.session_state.current_chat == chat_name:
                    st.session_state.current_chat = new_name
                st.session_state.edit_mode.pop(chat_name, None)
                st.session_state.new_chat_name = ""
                st.experimental_rerun()

# Display current chat and its history
st.subheader(f"Current Chat: {st.session_state.current_chat}")
for message in st.session_state.chats[st.session_state.current_chat]["messages"]:
    if message.startswith("You: "):
        st.markdown(f"""
        <div style="background-color: #f0f0f0; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
            <strong>🙋‍♂️ You:</strong> {message[5:]}
        </div>
        """, unsafe_allow_html=True)
    elif message.startswith("Bot: "):
        st.markdown(f"""
        <div style="background-color: #e6f3ff; padding: 10px; border-radius: 5px; margin-bottom: 10px;">
            <strong>🤖 Bot:</strong> {message[5:]}
        </div>
        """, unsafe_allow_html=True)

# File processing functions
def read_pdf(file):
    text = ""
    pdf_reader = PdfReader(file)
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

def read_docx(file):
    doc = DocxDocument(file)
    return "\n".join([para.text for para in doc.paragraphs])

def read_txt(file):
    return file.read().decode("utf-8")

def process_uploaded_file(uploaded_file):
    try:
        if uploaded_file.type == "application/pdf":
            text = read_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = read_docx(uploaded_file)
        elif uploaded_file.type == "text/plain":
            text = read_txt(uploaded_file)
        else:
            st.error("Unsupported file type.")
            return None
        
        return Document(text=text, metadata={"file_name": uploaded_file.name})
    except Exception as e:
        st.error(f"Error processing file: {e}")
        return None

# File uploader widget for multiple files
uploaded_files = st.file_uploader("Upload documents for current chat:", type=['txt', 'pdf', 'docx'], accept_multiple_files=True, key=f"uploader_{st.session_state.current_chat}")

if uploaded_files:
    # Process each uploaded file
    current_chat = st.session_state.chats[st.session_state.current_chat]
    new_documents = []
    for uploaded_file in uploaded_files:
        document = process_uploaded_file(uploaded_file)
        if document:
            new_documents.append(document)
            st.write(f"Loaded document: {uploaded_file.name}")
        else:
            st.write(f"Failed to load document: {uploaded_file.name}")
    
    if new_documents:
        # Add new documents to the current chat's document list
        current_chat["documents"].extend(new_documents)
        
        # Parse nodes and create/update the index
        parser = SimpleNodeParser()
        nodes = parser.get_nodes_from_documents(current_chat["documents"])
        current_chat["index"] = VectorStoreIndex(nodes, embed_model=embed_model)
        st.write(f"All documents processed and indexed for the current chat: {st.session_state.current_chat}")
    else:
        st.write("No valid documents were processed.")

# Display loaded documents for the current chat
st.subheader(f"Loaded Documents for Current Chat: {st.session_state.current_chat}")
for doc in st.session_state.chats[st.session_state.current_chat]["documents"]:
    st.write(f"- {doc.metadata['file_name']}: {doc.get_content()[:50]}...")  # Display file name and first 50 characters of each document

# Query input and processing
question = st.text_input("Enter your query:")

if st.button("Submit Query"):
    if question:
        current_chat = st.session_state.chats[st.session_state.current_chat]
        if not current_chat["documents"]:
            response = "No documents loaded for this chat. Please upload documents first."
        elif current_chat["index"] is None:
            response = "Index not created. Please try uploading documents again."
        else:             
            st.write("Processing query...")
            query_engine = current_chat["index"].as_query_engine(llm=llm)
            response = query_engine.query(question).response
        
        # Add question and response to current chat history
        current_chat["messages"].append(f"You: {question}")
        current_chat["messages"].append(f"Bot: {response}")
        
        # Rerun to update the display
        st.experimental_rerun()
    else:
        st.write("Please ask a question.")

# # Debug information
# st.sidebar.subheader("Debug Information")
# for chat_name, chat_data in st.session_state.chats.items():
#     st.sidebar.write(f"Chat: {chat_name}")
#     st.sidebar.write(f"  Documents: {len(chat_data['documents'])}")
#     st.sidebar.write(f"  Has Index: {'Yes' if chat_data['index'] is not None else 'No'}")
